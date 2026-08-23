"""TRAQ Word CV generator, independent from the web interface."""

from __future__ import annotations

from io import BytesIO
from typing import Any, Iterable, Mapping

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


FONT = "Calibri"
INK = RGBColor(0, 0, 0)


def _text(value: Any) -> str:
    return str(value or "").strip()


def _items(value: Any) -> list[Mapping[str, Any]]:
    if not isinstance(value, list):
        return []
    return [item for item in value if isinstance(item, Mapping)]


def _set_cell_border(paragraph: Any, color: str = "000000") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    old = p_pr.find(qn("w:pBdr"))
    if old is not None:
        p_pr.remove(old)
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    for key, value in (("val", "single"), ("sz", "5"), ("space", "1"), ("color", color)):
        bottom.set(qn(f"w:{key}"), value)
    borders.append(bottom)
    p_pr.append(borders)


def _configure_document(document: DocumentObject) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    # Values measured from GBENLE_PRAISE_CV.docx.
    section.top_margin = Inches(0.42)
    section.bottom_margin = Inches(0.35)
    section.left_margin = Inches(0.575)
    section.right_margin = Inches(0.575)

    normal = document.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0

    bullet = document.styles["List Bullet"]
    bullet.font.name = FONT
    bullet.font.size = Pt(11)
    bullet.paragraph_format.left_indent = Inches(0.22)
    bullet.paragraph_format.first_line_indent = Inches(-0.13)
    bullet.paragraph_format.space_after = Pt(0)

    if "CV Entry" not in document.styles:
        style = document.styles.add_style("CV Entry", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = normal
        style.paragraph_format.space_after = Pt(0)


def _section_heading(document: DocumentObject, title: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(3.2)
    paragraph.paragraph_format.space_after = Pt(1)
    run = paragraph.add_run(title.upper())
    run.bold = True
    run.font.name = FONT
    run.font.size = Pt(11)
    run.font.color.rgb = INK
    _set_cell_border(paragraph)


def _entry_header(document: DocumentObject, name: str, dates: str, *, italic: bool = False) -> None:
    paragraph = document.add_paragraph(style="CV Entry")
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(7.35), WD_TAB_ALIGNMENT.RIGHT)
    run = paragraph.add_run(name)
    run.bold = True
    run.italic = italic
    if dates:
        paragraph.add_run("\t")
        date_run = paragraph.add_run(dates)
        date_run.bold = True
        date_run.italic = italic


def _subtitle(document: DocumentObject, title: str, location: str = "", *, bold: bool = False) -> None:
    paragraph = document.add_paragraph(style="CV Entry")
    paragraph.paragraph_format.space_after = Pt(1)
    run = paragraph.add_run(title)
    run.italic = True
    run.bold = bold
    if location:
        suffix = paragraph.add_run(f" | {location}")
        suffix.italic = True
        suffix.bold = bold


def _bullets(document: DocumentObject, values: Iterable[Any]) -> None:
    for value in values:
        content = _text(value)
        if content:
            document.add_paragraph(content, style="List Bullet")


def _experience_section(document: DocumentObject, title: str, entries: Any) -> None:
    valid = [entry for entry in _items(entries) if _text(entry.get("organisation"))]
    if not valid:
        return
    _section_heading(document, title)
    for entry in valid:
        organisation = _text(entry.get("organisation"))
        location = _text(entry.get("location"))
        org_line = ", ".join(filter(None, [organisation, location]))
        org_paragraph = document.add_paragraph(style="CV Entry")
        org_paragraph.paragraph_format.space_before = Pt(0.6)
        org_run = org_paragraph.add_run(org_line)
        org_run.bold = True
        org_run.italic = True
        _entry_header(document, _text(entry.get("role")), _text(entry.get("dates")), italic=True)
        bullets = entry.get("bullets", [])
        _bullets(document, bullets if isinstance(bullets, list) else [])


def build_docx(data: dict[str, Any]) -> BytesIO:
    """Build a styled CV and return a seeked in-memory .docx file."""
    document = Document()
    _configure_document(document)

    name = _text(data.get("full_name")) or "CURRICULUM VITAE"
    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.space_after = Pt(2)
    run = heading.add_run(name.upper())
    run.bold = True
    run.font.name = FONT
    run.font.size = Pt(11)
    run.font.color.rgb = INK

    contact = " | ".join(filter(None, [_text(data.get("phone")), _text(data.get("email"))]))
    if contact:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(contact)

    linkedin = _text(data.get("linkedin"))
    if linkedin:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(3)
        link_run = paragraph.add_run(linkedin)
        link_run.underline = True
        link_run.font.color.rgb = RGBColor(0, 0, 255)

    education = [entry for entry in _items(data.get("education")) if _text(entry.get("school"))]
    if education:
        _section_heading(document, "Education")
        for entry in education:
            school_line = ", ".join(filter(None, [_text(entry.get("school")), _text(entry.get("location"))]))
            _entry_header(document, school_line, _text(entry.get("dates")))
            _subtitle(document, _text(entry.get("degree")), bold=True)
            detail = _text(entry.get("detail"))
            if detail:
                detail_paragraph = document.add_paragraph(style="CV Entry")
                detail_run = detail_paragraph.add_run(detail)
                detail_run.bold = True
                detail_run.italic = True

    _experience_section(document, "Work Experience", data.get("work_experience"))
    _experience_section(document, "Leadership & Volunteer Experience", data.get("leadership_experience"))

    certifications = [_text(item) for item in data.get("certifications", []) if _text(item)] if isinstance(data.get("certifications"), list) else []
    if certifications:
        _section_heading(document, "Certifications & Courses")
        _bullets(document, certifications)

    technical = _text(data.get("technical_skills"))
    strengths = _text(data.get("strengths"))
    if technical or strengths:
        _section_heading(document, "Skills & Interests")
        if technical:
            paragraph = document.add_paragraph(style="CV Entry")
            paragraph.add_run("Technical skills: ").bold = True
            paragraph.add_run(technical)
        if strengths:
            paragraph = document.add_paragraph(style="CV Entry")
            paragraph.add_run("Strengths: ").bold = True
            paragraph.add_run(strengths)

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output
